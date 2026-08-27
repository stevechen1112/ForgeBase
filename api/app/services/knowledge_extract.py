from __future__ import annotations

from io import BytesIO

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
INDEXABLE_MIME_TYPES = {"application/pdf", DOCX_MIME}


class NeedsOCR(Exception):
    """The file has no extractable text; do not pretend it was read."""


def is_indexable_document(mime_type: str | None, filename: str | None = None) -> bool:
    mime = (mime_type or "").lower()
    if mime in INDEXABLE_MIME_TYPES:
        return True
    suffix = (filename or "").lower()
    return suffix.endswith(".pdf") or suffix.endswith(".docx")


def extract_document_pages(data: bytes, mime_type: str, filename: str = "") -> list[tuple[int | None, str]]:
    mime = (mime_type or "").lower()
    name = (filename or "").lower()
    if mime == "application/pdf" or name.endswith(".pdf"):
        return _extract_pdf(data)
    if mime == DOCX_MIME or name.endswith(".docx"):
        return _extract_docx(data)
    raise NeedsOCR("unsupported document type")


def _extract_pdf(data: bytes) -> list[tuple[int | None, str]]:
    from pypdf import PdfReader

    reader = PdfReader(BytesIO(data))
    pages: list[tuple[int | None, str]] = []
    for index, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        if text:
            pages.append((index, text))
    if not pages:
        raise NeedsOCR("pdf_has_no_extractable_text")
    return pages


def _extract_docx(data: bytes) -> list[tuple[int | None, str]]:
    from docx import Document

    document = Document(BytesIO(data))
    paragraphs = [item.text.strip() for item in document.paragraphs if item.text and item.text.strip()]
    # Comments and revision markup are not part of document.paragraphs.
    text = "\n".join(paragraphs).strip()
    if not text:
        raise NeedsOCR("docx_has_no_extractable_text")
    return [(None, text)]
